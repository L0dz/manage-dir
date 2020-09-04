from tkinter import *
from tkinter import messagebox
from os import getcwd
from win32com import client
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from time import sleep
#####################
########TO DO########

#####################



class main_win:
    def WindowNotOpen(self):
        self.window_open = False

    def HandInUpdate(self, dal, al):
        #xx/xx/xxxx
        #0123456789
        #Wipe temporary list
        self.temporary_list = []

        #Translate date into int
        try:
            from_=int(dal[6:] + dal[3:5] + dal[:2])
            to_=int(al[6:] + al[3:5] + al[:2])
        except Exception as e:
            print(e)
            self.handin_window.destroy()
            self.window_open = False
            return

        #Check wheter each date is between the user input and fill the temporary list
        for element, el in self.current_list:
            check = int(element[3][6:] + element[3][3:5] + element[3][:2])
            if check >= from_ and check <= to_ and el != 'Requisito':
                #print(f'{check} is between {from_} and {to_}')
                self.temporary_list.append((element, 'Consegnato'))
            else:
                self.temporary_list.append((element, el))

        #Update actual list
        self.current_list = self.temporary_list
        self.temporary_list = []


        self.handin_window.destroy()
        self.window_open = False


        self.FilterUpdate()


    def HandIn(self):
        if not self.window_open:
            self.handin_window = Toplevel(self.main, bg=RADIO_BACKGROUND)
            self.window_open = True
        else:
            return

        self.handin_window.protocol('WM_DELETE_WINDOW', lambda: [self.handin_window.destroy(), self.WindowNotOpen()])
        
        #GUI
        Label(self.handin_window, text='Consegna dal', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND).grid(row=0, column=0)
        dal=Entry(self.handin_window)
        dal.grid(row=0, column=1)

        Label(self.handin_window, text='Al', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND ).grid(row=1, column=0)
        al=Entry(self.handin_window)
        al.grid(row=1, column=1)

        Button(self.handin_window, text='Fatto', bg = BUTTON_BACKGROUND, fg = BUTTON_FOREGROUND,  font=('Arial', 14), command= lambda: self.HandInUpdate(dal.get(), al.get())).grid(row=2, column=0)
        
        self.saved = False
    
    def EditList(self, new_person, index):
        #Delete old person
        self.list.delete(index)
        self.current_list.pop(index)

        #Add modified person
        self.current_list.append(new_person)
        self.list.insert(index, new_person[0])

        self.FilterUpdate()

        self.saved = False
        self.window_open = False
        self.edit_window.destroy()

    def Edit(self):
        #Create new window  
        if not self.window_open:
            self.edit_window = Toplevel(self.main, bg=RADIO_BACKGROUND)
            self.window_open = True
        else:
            return

        self.edit_window.protocol('WM_DELETE_WINDOW', lambda: [self.edit_window.destroy(), self.WindowNotOpen()])
        self.temporary_list = []

        #Get active person
        person = self.list.get(ANCHOR)
        if person == '':
            print('Nothing is selected')
            self.window_open = False
            self.edit_window.destroy()
            return

        for element in person:
            self.temporary_list.append(element)

        index = 0
        for el, el1 in self.current_list:
            #print(el, el1)
            if self.temporary_list == el:
                person = [el, el1]
                print(f'CHOOSEN = \n {el} ----> {el1}')
                break
            index+=1

        #print(person)
        #Display GUI
        Label(self.edit_window, text='Nome', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND,  font=('Arial', 12)).grid(row=0, column = 0)
        name = Entry(self.edit_window)
        name.grid(row=0, column = 1)
        name.insert(END, person[0][0])

        Label(self.edit_window, text='Cognome', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND,  font=('Arial', 12)).grid(row=1, column = 0)
        surname = Entry(self.edit_window)
        surname.grid(row=1, column = 1)
        surname.insert(END, person[0][1])

        Label(self.edit_window, text='Data di nascita', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND,  font=('Arial', 12)).grid(row=2, column = 0)
        born = Entry(self.edit_window)
        born.grid(row=2, column = 1)
        born.insert(END, person[0][2])

        Label(self.edit_window, text='Data di ingresso', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND,  font=('Arial', 12)).grid(row=3, column = 0)
        ing = Entry(self.edit_window)
        ing.grid(row=3, column = 1)
        ing.insert(END, person[0][3])

        Label(self.edit_window, text='Data di uscita', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND,  font=('Arial', 12)).grid(row=4, column = 0)
        usc = Entry(self.edit_window)
        usc.grid(row=4, column = 1)
        usc.insert(END, person[0][4])

        Label(self.edit_window, text='Diagnosi', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND,  font=('Arial', 12)).grid(row=5, column = 0)
        dia = Entry(self.edit_window)
        dia.grid(row=5, column = 1)
        dia.insert(END, person[0][5])

        Label(self.edit_window, text='Numero di cartella', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND, font=('Arial', 12)).grid(row=6, column = 0)
        dir = Entry(self.edit_window)
        dir.grid(row=6, column = 1)
        dir.insert(END, person[0][6])

        state = StringVar()
        state.set(person[1])
        stato=['Non Consegnato',
               'Consegnato',
               'Requisito']

        
        for word in stato:
            Radiobutton(self.edit_window, text=word, bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND, variable=state, value = word).grid(row=7, column=stato.index(word))


        Button(self.edit_window, text='Fatto', bg = BUTTON_BACKGROUND, fg = BUTTON_FOREGROUND,  font=('Arial',16), command = lambda: self.EditList([[name.get(), surname.get(), born.get(), ing.get(), usc.get(), dia.get(), dir.get()], state.get()], index)).grid(row=8)

        self.saved = False



    def Print(self):
        #Get text from current list
        text=''
        if self.choice == 'Tutto':
            for element, el2 in self.current_list:
                text+='-' + element[0] + ' ' + element[1] + ' ' + element[3] + ' ' + element[6] + ' ' +  el2 + '\n'
        else:
            for el in range(self.list.size()):
                text+='-' + self.list.get(el)[0] + ' '  + self.list.get(el)[1] + ' '  + self.list.get(el)[3] + ' '  + self.list.get(el)[6] + ' ' + self.choice + '\n'


        #Elenco delle cartelle %s:  (consegnate, non consegnate, requisite, '')
        #-Nome, cognome, data di ingresso, numero di cartella


        #Create new doc
        document = Document()
        document.add_paragraph(' ')
        document.save('list.docx')

        document = Document('list.docx')
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        #Add pic
        document.add_picture('asl.jpg',width=Inches(1.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        #Add head
        head = document.add_paragraph()
        font = head.add_run('U.O.C. CHIRURGIA VASCOLARE\nASL BRINDISI P.O. \'PERRINO\'\nDIRETTORE DOTT. GABRIELE MARITATI FEBVS').font
        font.size = Pt(18)
        head.alignment = 1

        #Add List
        body = document.add_paragraph()
        font = body.add_run(f'Elenco delle cartelle ({self.choice}):\n').font
        font.size = Pt(18)
        for element in range(self.list.size()):
            font = body.add_run('-' + self.list.get(element)[0] + '  ' + self.list.get(element)[1] +  '  ' + self.list.get(element)[3] +  '  ' + self.list.get(element)[4] + '  ' + self.list.get(element)[6] + '\n').font
            font.size = Pt(18)

        #Save doc
        document.save('list.docx')

        #Print doc on paper
        word = client.Dispatch("Word.Application")
        word.Documents.Open(getcwd() + '\\list.docx')
        word.ActiveDocument.PrintOut()
        sleep(2)
        word.ActiveDocument.Close()

    def FilterUpdate(self):

        if self.search:
            self.SearchClose()

        self.temporary_list = []

        #Filter the actual list
        #print(self.current_list)
        if self.choice == 'Tutto':
            for element, el2 in self.current_list:
                print('\n\n\n' + str(element))
                self.temporary_list.append(element)
        else:
            for element in range(len(self.current_list)):
                if self.current_list[element][1] == self.choice:
                    self.temporary_list.append(self.current_list[element][0])

        #Wipe actual list
        self.list.delete(0, END)

        #Make changes visible
        for element in self.temporary_list:
            self.list.insert(0, element)

        self.temporary_list = []

        if self.search:
            self.search = False
            self.SearchUpdate()

        self.OrderUpdate()


    def SearchClose(self):
        #Wipe list
        self.list.delete(0,END)

        #Restore original list
        for element, el2 in self.current_list:
            self.list.insert(0, element)

        self.x_active = False
        self.search_list = []

        self.OrderUpdate()

        try:
            self.close_search.destroy()
        except Exception as e:
            print(e)

    def SearchUpdate(self):
        #Get user's inputs
        self.keyword = self.word.get()
        self.type = self.searchtype.get()
        #print(str(self.keyword) + ' ' + str(self.type))

        #Get current list
        if not self.x_active:
            for element in range(self.list.size()):
                self.search_list.append(self.list.get(element))

        if self.type == 'Nome':
            newlist = [s for s in self.search_list if self.keyword.lower() in s[0].lower()]
        elif self.type == 'Cognome':
            newlist = [s for s in self.search_list if self.keyword.lower() in s[1].lower()]
        elif self.type == 'Numero di cartella':
            newlist = [s for s in self.search_list if self.keyword.lower() in s[6].lower()]
        elif self.type == 'Data di ingresso':
            newlist = [s for s in self.search_list if self.keyword.lower() in s[3].lower()]
        elif self.type == 'Diagnosi':
            newlist = [s for s in self.search_list if self.keyword.lower() in s[5].lower()]

        #Apply list change
        self.list.delete(0,END)

        for element in newlist:
            self.list.insert(0, element)


        self.OrderUpdate()
        self.saved = False
        


        #Close toplevel window
        self.search_window.destroy()
        self.window_open = False

        #Create X Button
        if not self.x_active:
            self.close_search = Button(self.main, text = 'X', bg = BUTTON_BACKGROUND, fg = BUTTON_FOREGROUND, font=('Arial', 16), command = self.SearchClose)#470 260
            self.close_search.place(x=700, y=260)
            self.x_active = True


    def OnClose(self):
        print(self.saved)
        if not self.saved:
            if messagebox.askyesno('', 'Vuoi uscire senza salvare?'):
                self.main.destroy()
        else:
            self.main.destroy()
    
    def Save(self):

        #Write in file the information
        text=''
        with open('pazienti.txt', 'w') as file:
            for element, hide in self.current_list:
                for piece in element:
                    text+=piece
                    text+='-'
                text+=hide
                text+='-\n'

            file.write(text)

        self.saved = True
        #Wipe temporary list
        self.temporary_list = []

    def Search(self):
        #Create new window
        if not self.window_open:
            self.search_window = Toplevel(self.main, bg=RADIO_BACKGROUND)
            self.window_open = True
        else:
            return

        self.search_window.protocol('WM_DELETE_WINDOW', lambda: [self.search_window.destroy(), self.WindowNotOpen()])
        #Create a text label and get an input word
        Label(self.search_window, text = 'Cerca', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND, font=('Arial', 13)).grid(row=0)
        self.word = Entry(self.search_window)
        self.word.grid(row=1)

        #Possible types
        search_list = [
            'Nome',
            'Cognome',
            'Numero di cartella',
            'Data di ingresso',
            'Diagnosi'
            ]

        #GUI
        self.searchtype = StringVar()
        self.searchtype.set('Numero di cartella')
        for types in search_list:
            Radiobutton(self.search_window, text=types, bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND, variable = self.searchtype, value = types).grid(row=search_list.index(types)+2)


        self.search = True

        Button(self.search_window, text='Fatto', bg = BUTTON_BACKGROUND, fg = BUTTON_FOREGROUND,  font=('Arial', 16), command = self.SearchUpdate).grid(row=search_list.index(types)+3)

        

    def ListUpdate(self):
        #Add in the list
        person = []
        try:
            self.people = (self.name.get(), self.surname.get(), self.born.get(), self.from_.get(), self.to_.get(), self.diagnosis.get(), self.dir_number.get())
            person.append([self.name.get(), self.surname.get(), self.born.get(), self.from_.get(), self.to_.get(), self.diagnosis.get(), self.dir_number.get()])
            #print(self.surname.get())
            self.list.insert(0, self.people)
        except Exception as e:
            print(e)

        try:
            self.people = (self.people, self.state.get())
            person.append(self.state.get())
        except Exception as e:
            print(e)

        #Destroy the window
        self.add_window.destroy()
        self.window_open = False

        #Update the order
        self.OrderUpdate()

        #Set saved false
        self.saved = False
        

        #Update parallel list
        self.current_list.append(person)

    def OrderUpdate(self):
        try:
            self.order = self.same.get()
        except Exception as e:
            print(e)
            print('same is null')
        #Get current list
        for element in range(self.list.size()):
            self.temporary_list.append(self.list.get(element))
            #print(self.list.get(element))

        #Delete list items
        self.list.delete(0,END)

        #print(self.current_list)
         
        try:
        #Sort current list
            if self.order == 'Nome':
                self.temporary_list = sorted(self.temporary_list, key = lambda x: x[0].lower())
                self.current_list = sorted(self.current_list, key = lambda x: x[0][0].lower())
            elif self.order == 'Cognome':
                self.temporary_list = sorted(self.temporary_list, key = lambda x: x[1].lower())
                self.current_list = sorted(self.current_list, key = lambda x: x[0][1].lower())
            elif self.order == 'Data di nascita':
                self.temporary_list = sorted(self.temporary_list, key = lambda x: x[2][6:] + x[2][3:5] + x[2][0:2])
                self.current_list = sorted(self.current_list, key = lambda x: x[0][2][6:] + x[0][2][3:5] + x[0][2][0:2])
            elif self.order == 'Numero di cartella':
                self.temporary_list = sorted(self.temporary_list, key = lambda x: x[6])
                self.current_list = sorted(self.current_list, key = lambda x: x[0][6])
            elif self.order == 'Data di ricovero':
                self.temporary_list = sorted(self.temporary_list, key = lambda x: x[3][6:] + x[3][3:5] + x[3][0:2])
                self.current_list = sorted(self.current_list, key = lambda x: x[0][3][6:] + x[0][3][3:5] + x[0][3][0:2])
        except Exception as e:
            print(e)

        #Update actual list
        for element in self.temporary_list:
            if len(element[1]) != 1:
                self.list.insert(END, element)


        #Reset temporary list
        self.temporary_list = []

        #self.saved = False
        try:
            self.order_window.destroy()
            self.window_open = False
        except Exception as e:
            print(e)


    def Add(self):
        #Create new window
        if not self.window_open:
            self.add_window = Toplevel(self.main, bg=RADIO_BACKGROUND)
            self.window_open = True
        else:
            return
        self.add_window.title('Add')
        self.add_window.protocol('WM_DELETE_WINDOW', lambda: [self.add_window.destroy(), self.WindowNotOpen()])

        #GUI
        Label(self.add_window, text='Nome', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND).grid(row=0, column=0)
        self.name = Entry(self.add_window)
        self.name.grid(row=0,column = 1)

        Label(self.add_window, text='Cognome', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND).grid(row=1, column=0)
        self.surname = Entry(self.add_window)
        self.surname.grid(row=1,column = 1)

        Label(self.add_window, text='Data di nascita', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND).grid(row=2, column=0)
        self.born = Entry(self.add_window)
        self.born.grid(row=2,column = 1)

        Label(self.add_window, text='Ricoverato dal', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND).grid(row=3, column=0)
        self.from_ = Entry(self.add_window)
        self.from_.grid(row=3,column = 1)

        Label(self.add_window, text='Al', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND).grid(row=3, column=2)
        self.to_ = Entry(self.add_window)
        self.to_.grid(row=3,column = 3)

        Label(self.add_window, text='Diagnosi', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND).grid(row=4, column=0)
        self.diagnosis = Entry(self.add_window)
        self.diagnosis.grid(row=4,column = 1)

        Label(self.add_window, text='Numero di cartella', bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND).grid(row=5, column=0)
        self.dir_number = Entry(self.add_window)
        self.dir_number.grid(row=5,column = 1)
        
        self.state = StringVar()
        self.state.set('Non Consegnato')
        stato = [
            'Non Consegnato',
            'Consegnato',
            'Requisito'
            ]

        x=0
        for word in stato:
            Radiobutton(self.add_window, text=word, bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND, variable=self.state, value=word).grid(row=6, column=x)
            x += 1
        
        #Done button
        done = Button(self.add_window, text='Fatto', bg = BUTTON_BACKGROUND, fg = BUTTON_FOREGROUND,  font=('Arial', 16), command = self.ListUpdate).grid(row=7)
        

    def Remove(self):
        #Remove selected item from the list
        ind = self.list.index(self.list.curselection())
        self.list.delete(ANCHOR)
        self.saved = False

        self.current_list.pop(ind)

    def Order(self):
        #Create new window
        if not self.window_open:
            self.order_window = Toplevel(self.main, bg=RADIO_BACKGROUND)
            self.window_open = True
        else:
            return

        self.order_window.protocol('WM_DELETE_WINDOW', lambda: [self.order_window.destroy(), self.WindowNotOpen()])

        #Assign values to each radio button
        order_list = [
            'Nome',
            'Cognome',
           'Data di nascita',
           'Numero di cartella',
           'Data di ricovero']

        self.same = StringVar()
        self.same.set('Numero di cartella')
        for orders in order_list:
            Radiobutton(self.order_window, text = orders, bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND, variable=self.same, value = orders).grid(row=order_list.index(orders))

        Button(self.order_window, text='Ordina', bg = BUTTON_BACKGROUND, fg = BUTTON_FOREGROUND,  font=('Arial', 16), command = self.OrderUpdate).grid(row = len(order_list)+1)

        self.saved = False


    def ManageFilter(self, arg):
        self.choice = arg
        self.FilterUpdate()

    def __init__(self):
        self.main = Tk()
        self.main.geometry('1000x800')
        self.main.title('Made by Andrea Testini')
        self.main.resizable(False,False)
        self.main.protocol('WM_DELETE_WINDOW', self.OnClose)
        
        self.main.config(bg=RADIO_BACKGROUND)

        self.scroll = Scrollbar(self.main)
        self.scroll.pack(side=RIGHT, fill=Y)

        self.list = Listbox(self.main, bg = BUTTON_BACKGROUND, fg = RADIO_FOREGROUND, font=('Verdana, 16'))#, bg='#BFBFBF')
        self.list.config(yscrollcommand=self.scroll.set)
        self.scroll.config(command=self.list.yview_scroll)
        self.people = []
        self.list.pack(fill = X)
        self.current_list = []
        self.temporary_list = []
        self.search_list = []
        self.choice = 'Tutto'


        self.order_window = None
        self.order = 'Cognome'
        self.same = None
        
        self.window_open = False
        self.saved = True
        self.x_active = False
        self.search = False

        #ADD READING EXISTING LIST FROM A FILE
        try:
            with open('pazienti.txt', 'r') as file:
                text = file.read()
        except Exception as e:
            print('Error: ' + str(e))
            with open('pazienti.txt', 'w') as file:
                pass
        
        try:
            with open('pazienti.txt') as file:
                word = ''
                hidden = ''
                count=0
                tup=[]
                for letters in text:
                    if letters == '\n':
                        pass
                    elif letters != '-':
                        word += letters
                    else:
                        count += 1
                        if count > 7:
                            hidden = word
                        else:
                            tup.append(word)
                        
                        #print(tup)
                        if count == 8:
                            count = 0
                            self.list.insert(0, tup)
                            self.current_list.append((tup, hidden))
                            tup=[]

                        word = ''
                self.OrderUpdate()
        except Exception as e:
            print(e)
            print('File hasn\'t been read')

        #Main GUI
        Button(self.main, text='Aggiungi', bg = BUTTON_BACKGROUND, fg = BUTTON_FOREGROUND, font=('Arial', 16), command = self.Add).place(x=10, y=260)
        Button(self.main, text='Rimuovi', bg = BUTTON_BACKGROUND, fg = BUTTON_FOREGROUND, font=('Arial', 16), command = self.Remove).place(x=120, y=260)
        Button(self.main, text='Ordina per...', bg = BUTTON_BACKGROUND, fg = BUTTON_FOREGROUND, font=('Arial', 16), command = self.Order).place(x=230, y=260)
        Button(self.main, text='Cerca', bg = BUTTON_BACKGROUND, fg = BUTTON_FOREGROUND, font=('Arial', 16), command = self.Search).place(x=380, y=260)
        Button(self.main, text='Salva', bg = BUTTON_BACKGROUND, fg = BUTTON_FOREGROUND, font=('Arial', 16), command = self.Save).place(x=470, y=260)
        Button(self.main, text='Stampa', bg = BUTTON_BACKGROUND, fg = BUTTON_FOREGROUND, font=('Arial', 16), command = self.Print).place(x=10, y=310)
        Button(self.main, text='Modifica', bg = BUTTON_BACKGROUND, fg = BUTTON_FOREGROUND, font=('Arial', 16), command = self.Edit).place(x=120, y=310)
        Button(self.main, text='Consegna', bg = BUTTON_BACKGROUND, fg = BUTTON_FOREGROUND, font=('Arial', 16), command = self.HandIn).place(x=230, y=310)

        self.consegnato = StringVar()
        self.consegnato.set('Tutto')
        stato = [
            'Tutto',
            'Non Consegnato',
            'Consegnato',
            'Requisito'
            ]
        Radiobutton(self.main, text=stato[0], bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND, variable=self.consegnato, value=stato[0], command = lambda: self.ManageFilter(stato[0])).place(x=560, y=260)
        Radiobutton(self.main, text=stato[1], bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND, variable=self.consegnato, value=stato[1], command = lambda: self.ManageFilter(stato[1])).place(x=560, y=285)
        Radiobutton(self.main, text=stato[2], bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND, variable=self.consegnato, value=stato[2], command = lambda: self.ManageFilter(stato[2])).place(x=560, y=310)
        Radiobutton(self.main, text=stato[3], bg = RADIO_BACKGROUND, fg = RADIO_FOREGROUND, variable=self.consegnato, value=stato[3], command = lambda: self.ManageFilter(stato[3])).place(x=560, y=335)

        #Button(self.main, text='Aggiorna', font=('Arial', 16), command=self.FilterUpdate).place(x=590, y=ypos+20)
        self.main.mainloop()
        
BUTTON_BACKGROUND = '#0D0D0D'
BUTTON_FOREGROUND = '#A6A6A6'
RADIO_BACKGROUND = '#262626'
RADIO_FOREGROUND = '#737373'
window = main_win()